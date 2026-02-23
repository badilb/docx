import os
import math
import subprocess
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import qn
from docx.oxml import OxmlElement


class DocxService:
    def __init__(self, output_dir: str):
        self.output_dir = output_dir

    @staticmethod
    def _section_has_footer(section) -> bool:
        footer = section.footer
        for p in footer.paragraphs:
            if p.text.strip():
                return True
        return bool(footer.tables)

    @staticmethod
    def _clear_footer(footer):
        fp = footer._element
        for child in list(fp):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("tbl", "p", "sdt"):
                fp.remove(child)

    def _build_footer_image(self, qr_code_path: str, logo_path: str | None, out_path: str):
        DPI = 300
        # Размер в дюймах → пикселях при 300dpi
        W_in, H_in = 3.2, 0.75
        W = int(W_in * DPI)
        H = int(H_in * DPI)

        PAD = int(0.08 * DPI)      # отступ контента от рамки
        RADIUS = int(0.12 * DPI)   # скругление углов
        BORDER = 2                 # толщина линии рамки в px
        COLOR = (50, 110, 190)
        DASH = int(0.09 * DPI)
        GAP  = int(0.05 * DPI)

        img = Image.new("RGB", (W, H), (255, 255, 255))
        draw = ImageDraw.Draw(img)

        # --- Белый фон со скруглёнными углами ---
        def fill_rounded(draw, x0, y0, x1, y1, r, color):
            draw.rectangle([x0+r, y0, x1-r, y1], fill=color)
            draw.rectangle([x0, y0+r, x1, y1-r], fill=color)
            for cx, cy in [(x0+r, y0+r), (x1-r, y0+r), (x0+r, y1-r), (x1-r, y1-r)]:
                draw.ellipse([cx-r, cy-r, cx+r, cy+r], fill=color)

        fill_rounded(draw, 0, 0, W-1, H-1, RADIUS, (255, 255, 255))

        # --- Пунктирная рамка ---
        def dashed_line(draw, p1, p2, dash, gap, color, width):
            x1, y1 = p1; x2, y2 = p2
            length = math.hypot(x2-x1, y2-y1)
            if length == 0: return
            dx, dy = (x2-x1)/length, (y2-y1)/length
            pos, on = 0, True
            while pos < length:
                end = min(pos + (dash if on else gap), length)
                if on:
                    draw.line(
                        [(int(x1+dx*pos), int(y1+dy*pos)), (int(x1+dx*end), int(y1+dy*end))],
                        fill=color, width=width
                    )
                pos, on = end, not on

        def dashed_arc(draw, cx, cy, r, a_start, a_end, dash_deg, gap_deg, color, width):
            a, on = a_start, True
            while a < a_end:
                seg = min(a + (dash_deg if on else gap_deg), a_end)
                if on:
                    pts = [(int(cx + r*math.cos(math.radians(t))),
                            int(cy + r*math.sin(math.radians(t))))
                           for t in range(int(a), int(seg)+1, 2)]
                    if len(pts) >= 2:
                        draw.line(pts, fill=color, width=width)
                a, on = seg, not on

        b = BORDER
        r = RADIUS
        x0, y0, x1, y1 = b, b, W-b-1, H-b-1
        D, G = DASH, GAP
        DA, GA = 20, 12  # градусы для дуг

        dashed_line(draw, (x0+r, y0), (x1-r, y0), D, G, COLOR, b)
        dashed_line(draw, (x0+r, y1), (x1-r, y1), D, G, COLOR, b)
        dashed_line(draw, (x0, y0+r), (x0, y1-r), D, G, COLOR, b)
        dashed_line(draw, (x1, y0+r), (x1, y1-r), D, G, COLOR, b)
        dashed_arc(draw, x0+r, y0+r, r, 180, 270, DA, GA, COLOR, b)
        dashed_arc(draw, x1-r, y0+r, r, 270, 360, DA, GA, COLOR, b)
        dashed_arc(draw, x1-r, y1-r, r,   0,  90, DA, GA, COLOR, b)
        dashed_arc(draw, x0+r, y1-r, r,  90, 180, DA, GA, COLOR, b)

        # --- QR-код ---
        QR_SIZE = H - PAD*2
        qr = Image.open(qr_code_path).convert("RGBA").resize((QR_SIZE, QR_SIZE), Image.LANCZOS)
        qr_x = PAD + int(0.02*DPI)
        qr_y = PAD
        bg = Image.new("RGB", (QR_SIZE, QR_SIZE), (255, 255, 255))
        bg.paste(qr, mask=qr.split()[3])
        img.paste(bg, (qr_x, qr_y))

        # --- Шрифт ---
        font_size = int(0.13 * DPI)  # ~11pt при 300dpi
        try:
            fp = "C:/Windows/Fonts/arial.ttf"
            if not os.path.exists(fp):
                fp = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
            font = ImageFont.truetype(fp, font_size)
        except Exception:
            font = ImageFont.load_default()

        # --- Текст ---
        line1 = "Подписи ЭЦП проверены НУЦ РК"
        line2 = "Документ подписан в сервис "
        text_color = (0, 0, 0)

        lh = font.getbbox("A")[3]  # высота строки
        gap_lines = int(0.03 * DPI)
        total_h = lh * 2 + gap_lines
        text_x = qr_x + QR_SIZE + int(0.1 * DPI)
        text_y = (H - total_h) // 2

        draw.text((text_x, text_y), line1, font=font, fill=text_color)
        line2_y = text_y + lh + gap_lines
        draw.text((text_x, line2_y), line2, font=font, fill=text_color)

        # --- Логотип ---
        if logo_path and os.path.exists(logo_path):
            logo_h = int(lh * 1.5)
            logo = Image.open(logo_path).convert("RGBA")
            logo_w = int(logo.width * logo_h / logo.height)
            logo = logo.resize((logo_w, logo_h), Image.LANCZOS)
            logo_x = text_x + int(draw.textlength(line2, font=font))
            logo_y = line2_y + (lh - logo_h) // 2
            img.paste(logo, (logo_x, logo_y), logo)

        img.save(out_path, "PNG", dpi=(DPI, DPI))
        return W_in, H_in  # возвращаем дюймы

    def _add_qr_footer(self, footer, qr_code_path: str, logo_path: str | None = None):
        if not footer.paragraphs:
            footer._element.append(OxmlElement("w:p"))

        for p in footer.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            pPr = p._element.get_or_add_pPr()
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "right")
            old_jc = pPr.find(qn("w:jc"))
            if old_jc is not None:
                pPr.remove(old_jc)
            pPr.append(jc)

        img_path = os.path.join(self.output_dir, "_footer_img.png")
        w_in, h_in = self._build_footer_image(qr_code_path, logo_path, img_path)

        p = footer.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.add_run().add_picture(img_path, width=Inches(w_in))

    def process_docx(self, template_path: str, output_docx_path: str,
                     qr_code_path: str, logo_path: str | None = None) -> str:
        doc = Document(template_path)
        for section in doc.sections:
            section.different_first_page_header_footer = False
            section.footer_distance = Cm(0.3)
            has_footer = self._section_has_footer(section)
            sectPr = section._sectPr
            for ref in sectPr.findall(qn("w:footerReference")):
                sectPr.remove(ref)
            self._clear_footer(section.footer)
            self._add_qr_footer(section.footer, qr_code_path, logo_path)
            print(f"[DocxService] Футер {'заменён' if has_footer else 'добавлен'}.")
        doc.save(output_docx_path)
        return output_docx_path

    def convert_to_pdf(self, docx_path: str) -> str:
        soffice = r"C:\Program Files\LibreOffice\program\soffice.exe"
        if not os.path.exists(soffice):
            soffice = "libreoffice"
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf",
             "--outdir", self.output_dir, docx_path],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice ошибка:\n{result.stderr}")
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(self.output_dir, f"{base_name}.pdf")
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF не найден: {pdf_path}")
        return pdf_path

    def process_and_convert(self, template_path: str, qr_code_path: str,
                            logo_path: str | None = None) -> str:
        base_name = os.path.splitext(os.path.basename(template_path))[0]
        tmp_docx = os.path.join(self.output_dir, f"_tmp_{base_name}.docx")
        self.process_docx(template_path, tmp_docx, qr_code_path, logo_path)
        pdf_path = self.convert_to_pdf(tmp_docx)
        try:
            os.remove(tmp_docx)
            os.remove(os.path.join(self.output_dir, "_footer_img.png"))
        except OSError:
            pass
        return pdf_path